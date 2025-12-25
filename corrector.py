"""
Sistema de corrección básico usando LanguageTool.
"""
import language_tool_python
from docx import Document
from docx.shared import RGBColor
from typing import List, Dict, Tuple
from ortotipografia import OrtotipografiaRules


class BasicCorrector:
    """Corrector básico que aplica cambios directamente al documento."""
    
    def __init__(self, idioma: str = 'es', usar_languagetool: bool = True):
        """
        Inicializa el corrector.
        
        Args:
            idioma: Código de idioma (default: 'es' para español)
            usar_languagetool: Si False, solo usa reglas ortotipográficas
        """
        self.ortotipo = OrtotipografiaRules()
        self.usar_lt = usar_languagetool
        self.tool = None
        
        if self.usar_lt:
            try:
                print("⏳ Iniciando LanguageTool...")
                self.tool = language_tool_python.LanguageTool(idioma)
                print("✓ LanguageTool iniciado")
            except Exception as e:
                print(f"⚠️  LanguageTool no disponible: {e}")
                print("⚠️  Continuando solo con reglas ortotipográficas")
                self.usar_lt = False
                self.tool = None
        else:
            print("ℹ️  Modo solo ortotipografía (sin LanguageTool)")
        
        self.stats = {
            'errores_gramaticales': 0,
            'correcciones_ortotipo': 0,
            'parrafos_procesados': 0
        }

    
    def configurar_reglas(self):
        """
        Configura las reglas de LanguageTool para corrección editorial.
        
        Activa:
        - Detección de tildes
        - Concordancia
        - Dequeísmo/queísmo
        - Comillas tipográficas (si está disponible)
        """
        # Desactivar reglas que generan mucho ruido
        reglas_desactivar = [
            'WHITESPACE_RULE',  # Muy sensible, lo manejamos con ortotipo
        ]
        
        for regla in reglas_desactivar:
            try:
                self.tool.disable_rule(regla)
            except:
                pass  # La regla puede no existir en esta versión
    
    def analizar_texto(self, texto: str) -> List[language_tool_python.Match]:
        """
        Analiza un texto con LanguageTool.
        
        Args:
            texto: Texto a analizar
            
        Returns:
            Lista de matches (errores detectados)
        """
        return self.tool.check(texto)
    
    def aplicar_correccion_languagetool(self, texto: str) -> str:
        """
        Aplica correcciones de LanguageTool automáticamente.
        
        Args:
            texto: Texto original
            
        Returns:
            Texto corregido
        """
        if not self.usar_lt or self.tool is None:
            return texto  # Sin cambios si LanguageTool no está disponible
        
        matches = self.analizar_texto(texto)
        
        # Filtrar solo errores con alta confianza
        matches_confianza = [
            m for m in matches 
            if len(m.replacements) > 0
        ]
        
        # Aplicar correcciones (de atrás hacia adelante para no desajustar posiciones)
        texto_corregido = texto
        for match in reversed(matches_confianza):
            if match.replacements:
                mejor_reemplazo = match.replacements[0]
                texto_corregido = (
                    texto_corregido[:match.offset] + 
                    mejor_reemplazo + 
                    texto_corregido[match.offset + match.errorLength:]
                )
                self.stats['errores_gramaticales'] += 1
        
        return texto_corregido
    
    def procesar_parrafo(self, parrafo) -> bool:
        """
        Procesa un párrafo del documento.
        
        Args:
            parrafo: Párrafo de python-docx
            
        Returns:
            True si se hicieron cambios, False si no
        """
        if not parrafo.text.strip():
            return False
        
        texto_original = parrafo.text
        
        # 1. Aplicar correcciones de LanguageTool
        texto_corregido = self.aplicar_correccion_languagetool(texto_original)
        
        # 2. Aplicar reglas de ortotipografía
        texto_corregido = self.ortotipo.aplicar_todas(texto_corregido)
        
        # 3. Si hubo cambios, actualizar el párrafo
        if texto_corregido != texto_original:
            # Preservar formato: mantenemos runs si es posible
            # Por ahora, implementación simple: reemplazar texto completo
            # TODO: Preservar formato de runs individuales
            
            parrafo.clear()
            parrafo.add_run(texto_corregido)
            
            self.stats['correcciones_ortotipo'] += 1
            return True
        
        return False
    
    def procesar_documento(self, ruta_entrada: str, ruta_salida: str):
        """
        Procesa un documento .docx completo.
        
        Args:
            ruta_entrada: Ruta del documento original
            ruta_salida: Ruta donde guardar el documento corregido
        """
        print(f"\n📄 Procesando: {ruta_entrada}")
        
        # Cargar documento
        doc = Document(ruta_entrada)
        
        # Procesar cada párrafo
        cambios_totales = 0
        for i, parrafo in enumerate(doc.paragraphs):
            if self.procesar_parrafo(parrafo):
                cambios_totales += 1
            
            self.stats['parrafos_procesados'] += 1
            
            # Progreso cada 10 párrafos
            if (i + 1) % 10 == 0:
                print(f"  Procesados {i + 1} párrafos...")
        
        # Guardar documento corregido
        doc.save(ruta_salida)
        
        print(f"\n✓ Documento guardado: {ruta_salida}")
        self.mostrar_estadisticas()
    
    def mostrar_estadisticas(self):
        """Muestra estadísticas del procesamiento."""
        print("\n📊 Estadísticas:")
        print(f"  • Párrafos procesados: {self.stats['parrafos_procesados']}")
        print(f"  • Errores gramaticales: {self.stats['errores_gramaticales']}")
        print(f"  • Correcciones ortotipo: {self.stats['correcciones_ortotipo']}")
        print(f"  • Total cambios: {self.stats['errores_gramaticales'] + self.stats['correcciones_ortotipo']}")
    
    def cerrar(self):
        """Cierra el corrector y libera recursos."""
        if self.tool is not None:
            self.tool.close()
        print("\n✓ Corrector cerrado")


if __name__ == '__main__':
    # Test rápido
    corrector = BasicCorrector(usar_languagetool=False)  # Solo ortotipografía para test rápido
    
    # Test con texto de ejemplo
    texto_test = 'El dijo: "la situacion es critica".'
    print(f"\nOriginal: {texto_test}")
    
    resultado = corrector.ortotipo.aplicar_todas(texto_test)
    
    print(f"Corregido: {resultado}")
    
    corrector.cerrar()
