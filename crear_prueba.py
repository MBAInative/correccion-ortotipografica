"""
Script para crear un documento de prueba con errores típicos.
"""
from docx import Document
from docx.shared import Pt, RGBColor


def crear_documento_prueba(ruta_salida: str):
    """
    Crea un documento .docx con errores ortotipográficos y gramaticales típicos.
    """
    doc = Document()
    
    # Título
    titulo = doc.add_heading('Documento de Prueba - Corrección Ortotipográfica', level=1)
    
    # Párrafo 1: Comillas incorrectas
    p1 = doc.add_paragraph()
    p1.add_run('El portavoz declaró: "La situación es \'crítica\' según el informe".')
    
    # Párrafo 2: Rayas vs guiones
    p2 = doc.add_paragraph()
    p2.add_run('- Hola - dijo María.')
    
    # Párrafo 3: Espacios con símbolos
    p3 = doc.add_paragraph()
    p3.add_run('El crecimiento fue del 25 % en 2024. Costó 1500 € aproximadamente.')
    
    # Párrafo 4: Puntuación con comillas
    p4 = doc.add_paragraph()
    p4.add_run('Como dice el refrán: "a quien madruga, Dios le ayuda".')
    
    # Párrafo 5: Tildes faltantes (error gramatical)
    p5 = doc.add_paragraph()
    p5.add_run('El niño comio rapidamente porque tenia mucha hambre.')
    
    # Párrafo 6: Concordancia (error que LanguageTool debería detectar)
    p6 = doc.add_paragraph()
    p6.add_run('Los estudiante llegó tarde a la clase.')
    
    # Párrafo 7: Texto con varios errores combinados
    p7 = doc.add_paragraph()
    p7.add_run('- No se que hacer - dijo Pedro. "La solucion no es facil".')
    
    # Párrafo 8: Uso de unidades
    p8 = doc.add_paragraph()
    p8.add_run('La distancia es de 10 km y el peso de 5 kg.')
    
    # Guardar
    doc.save(ruta_salida)
    print(f"✓ Documento de prueba creado: {ruta_salida}")
    print("\nErrores incluidos:")
    print("  • Comillas inglesas en lugar de latinas")
    print("  • Guiones en lugar de rayas")
    print("  • Espacios normales en lugar de duros")
    print("  • Puntuación antes de comillas")
    print("  • Tildes faltantes")
    print("  • Errores de concordancia")


if __name__ == '__main__':
    crear_documento_prueba('documento_prueba.docx')
